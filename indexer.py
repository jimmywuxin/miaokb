"""kb 索引器：扫知识库 → 抽内容 → 建索引（SQLite + 简单倒排）

不做语义向量（demo 阶段先验证「整 + 联」的实际价值，词袋 + 词频就够看出效果）。
"""
from __future__ import annotations
import os
import re
import json
import sys
import sqlite3
import hashlib
from pathlib import Path
from datetime import datetime
from collections import Counter

# 知识库根目录：实际生效路径由 resolve_kb_root() 在运行时决定。
# 优先级：环境变量 KB_ROOT > 第一个 CLI 参数 > DEFAULT_KB_ROOT
# 这样同一份代码可以跨 macOS / Windows / Linux 使用，不用改源码。
DEFAULT_KB_ROOT = Path("/Volumes/mac mini outside/知识库")
KB_ROOT = DEFAULT_KB_ROOT
INDEX_DIR = Path.home() / ".miaokb"
DB_PATH = INDEX_DIR / "kb.db"

# 内置支持的文本类型（无依赖）
TEXT_EXT = {".md", ".txt"}
# 跳过这些
SKIP_DIRS = {".DS_Store", "__pycache__"}

# 可选依赖的扩展器，按文件后缀注册。未装对应库时静默跳过该类型。
EXT_EXTRACTORS: dict = {}


def _register_extractors() -> None:
    """惰性注册 docx/pdf/xlsx 提取器。缺哪个库就跳过哪个类型，不影响其他。"""
    try:
        from docx import Document  # type: ignore

        def extract_docx(path: Path) -> str:
            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs if p.text]
            # 表格里的文字也别漏（很多人把内容塞在表格里）
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
            return "\n".join(parts)

        EXT_EXTRACTORS[".docx"] = extract_docx
    except ImportError:
        pass

    try:
        from pypdf import PdfReader  # type: ignore

        def extract_pdf(path: Path) -> str:
            reader = PdfReader(str(path))
            parts: list = []
            for page in reader.pages:
                try:
                    parts.append(page.extract_text() or "")
                except Exception:
                    # 单页解析失败不阻塞整本 PDF
                    parts.append("")
            return "\n".join(parts)

        EXT_EXTRACTORS[".pdf"] = extract_pdf
    except ImportError:
        pass

    try:
        from openpyxl import load_workbook  # type: ignore

        def extract_xlsx(path: Path) -> str:
            wb = load_workbook(str(path), read_only=True, data_only=True)
            parts: list = []
            try:
                for ws in wb.worksheets:
                    parts.append(f"# {ws.title}")
                    for row in ws.iter_rows(values_only=True):
                        line = " ".join(
                            str(c) for c in row if c is not None
                        ).strip()
                        if line:
                            parts.append(line)
            finally:
                wb.close()
            return "\n".join(parts)

        EXT_EXTRACTORS[".xlsx"] = extract_xlsx
    except ImportError:
        pass


_register_extractors()
SUPPORTED_EXT = TEXT_EXT | set(EXT_EXTRACTORS.keys())


def resolve_kb_root(argv=None) -> Path:
    """按 env > CLI > default 顺序解析知识库根目录。

    用法示例：
      export KB_ROOT=/path/to/kb && python3 indexer.py
      python3 indexer.py /path/to/kb
      python3 indexer.py            # 用 DEFAULT_KB_ROOT
    """
    if argv is None:
        argv = sys.argv[1:]
    env = os.environ.get("KB_ROOT")
    if env:
        return Path(env)
    if argv and not argv[0].startswith("-"):
        return Path(argv[0])
    return DEFAULT_KB_ROOT


def extract_text(path: Path) -> str:
    """按扩展名分发到对应提取器，失败返回空串（不中断整个索引过程）"""
    ext = path.suffix.lower()
    if ext in TEXT_EXT:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""
    if ext in EXT_EXTRACTORS:
        try:
            return EXT_EXTRACTORS[ext](path)
        except Exception:
            return ""
    return ""


def tokenize(text: str) -> list[str]:
    """中文：保留整块 + 2-gram；英文：单词"""
    # 去掉 markdown 标记、html 标签
    text = re.sub(r"[#*`>_\-\[\]()<>|]", " ", text)
    text = re.sub(r"https?://\S+", " ", text)
    tokens = []
    for m in re.finditer(r"[\u4e00-\u9fff]+|[a-zA-Z]+", text):
        seg = m.group()
        if seg[0].isascii():
            if len(seg) >= 2:
                tokens.append(seg.lower())
        else:
            # 中文：整块 + 2-gram 必加；3-gram 选加（命中「灵龟公园」这种短语）
            if len(seg) >= 2:
                tokens.append(seg)
            if len(seg) >= 3:
                for i in range(len(seg) - 1):
                    tokens.append(seg[i:i+2])
            if len(seg) >= 4:
                # 3-gram 加权（出现次数 = 2，提升相对权重）
                for i in range(len(seg) - 2):
                    tokens.append(seg[i:i+3])
                    tokens.append(seg[i:i+3])  # 双重权重
    return tokens


def file_hash(content: str) -> str:
    return hashlib.md5(content.encode("utf-8")).hexdigest()[:12]


def build_index(root=None) -> dict:
    """增量建索引：对比老快照与这一轮 rglob 结果，只处理有变化的文件。

    - mtime 一致 → fast path 跳过(连 hash 都不用算)
    - mtime 变了但 hash 没变(被 touch)→ 仅更新 mtime,沿用旧 terms
    - 真改了/新增了 → 重抽内容+重写该文件的 terms
    - 老有、新没有 → 删除对应 files + terms

    首次跑(db 为空)老快照为空,所有文件走 INSERT 分支,等价全量重建。
    """
    import time
    if root is None:
        root = KB_ROOT
    INDEX_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    # 幂等创建表,不再 DROP —— v0.2.1 起走增量
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS files (
            id INTEGER PRIMARY KEY,
            path TEXT UNIQUE,
            rel_path TEXT,
            size INTEGER,
            mtime TEXT,
            content_hash TEXT,
            token_count INTEGER,
            summary TEXT
        );
        CREATE TABLE IF NOT EXISTS terms (
            term TEXT,
            file_id INTEGER,
            freq INTEGER,
            PRIMARY KEY (term, file_id)
        );
        CREATE INDEX IF NOT EXISTS idx_terms_term ON terms(term);
    """)

    # 老快照:path -> (file_id, mtime, content_hash)
    cur.execute("SELECT id, path, mtime, content_hash FROM files")
    old_index = {p: (fid, m, h) for fid, p, m, h in cur.fetchall()}

    t0 = time.perf_counter()
    new_paths = set()
    stats = {"added": 0, "modified": 0, "deleted": 0, "skipped": 0, "unsupported": 0}

    for path in root.rglob("*"):
        if not path.is_file():
            continue
        if any(part.startswith(".") for part in path.parts):
            continue
        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXT:
            stats["unsupported"] += 1
            continue

        path_str = str(path)
        new_paths.add(path_str)
        stat = path.stat()
        mtime_str = datetime.fromtimestamp(stat.st_mtime).isoformat()

        # fast path: mtime 一致 → 零成本跳过
        if path_str in old_index and old_index[path_str][1] == mtime_str:
            stats["skipped"] += 1
            continue

        # mtime 变了 → 抽内容、算 hash
        content = extract_text(path)
        if not content.strip():
            continue

        full_text = path.name + "\n" + content
        tokens = tokenize(full_text)
        if not tokens:
            continue

        h = file_hash(content)
        # hash 没变(被 touch 但内容没变) → 仅更新 mtime 即可
        if path_str in old_index and old_index[path_str][2] == h:
            fid = old_index[path_str][0]
            cur.execute("UPDATE files SET mtime=? WHERE id=?", (mtime_str, fid))
            stats["skipped"] += 1
            continue

        summary = re.sub(r"\s+", " ", content)[:200].strip()
        rel_path = str(path.relative_to(root))

        # UPSERT:已有就更新,没有就插入
        cur.execute("SELECT id FROM files WHERE path=?", (path_str,))
        row = cur.fetchone()
        if row:
            file_id = row[0]
            cur.execute(
                "UPDATE files SET rel_path=?, size=?, mtime=?, content_hash=?, token_count=?, summary=? WHERE id=?",
                (rel_path, stat.st_size, mtime_str, h, len(tokens), summary, file_id),
            )
        else:
            cur.execute(
                "INSERT INTO files (path, rel_path, size, mtime, content_hash, token_count, summary) "
                "VALUES (?, ?, ?, ?, ?, ?, ?)",
                (path_str, rel_path, stat.st_size, mtime_str, h, len(tokens), summary),
            )
            file_id = cur.lastrowid

        # 清旧 terms,写新
        cur.execute("DELETE FROM terms WHERE file_id=?", (file_id,))
        freq = Counter(tokens)
        # 过滤停用词 + 太低频（仅出现 1 次的扔掉省空间）
        STOP = {
            "的", "了", "和", "是", "在", "我", "有", "不", "这", "也", "就", "都",
            "而", "及", "或", "与", "等", "中", "为", "对", "可", "上", "下", "一个",
            "the", "a", "an", "of", "to", "in", "is", "it", "and", "or", "for", "on",
        }
        for term, f in freq.items():
            if term in STOP or f < 2 or len(term) < 2:
                continue
            cur.execute(
                "INSERT INTO terms (term, file_id, freq) VALUES (?, ?, ?)",
                (term, file_id, f),
            )

        if path_str in old_index:
            stats["modified"] += 1
        else:
            stats["added"] += 1

    # 删除检测:老有、这一轮没扫到的(改名/移走/真删了)
    for p in set(old_index.keys()) - new_paths:
        fid = old_index[p][0]
        cur.execute("DELETE FROM terms WHERE file_id=?", (fid,))
        cur.execute("DELETE FROM files WHERE id=?", (fid,))
        stats["deleted"] += 1

    conn.commit()
    conn.close()
    stats["elapsed_s"] = round(time.perf_counter() - t0, 2)
    stats["is_first_run"] = len(old_index) == 0
    return stats


if __name__ == "__main__":
    KB_ROOT = resolve_kb_root()
    print(f"开始索引…  知识库: {KB_ROOT}")
    result = build_index()
    label = "首次全量" if result["is_first_run"] else "增量"
    print(f"✅ 索引完成({label} {result['elapsed_s']}s)")
    print(
        f"   新增: {result['added']}  修改: {result['modified']}  "
        f"删除: {result['deleted']}  跳过: {result['skipped']}"
    )
    print(f"   不支持类型跳过: {result['unsupported']}")
    print(f"   索引位置: {DB_PATH}")
